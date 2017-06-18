from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

page1=["Vehicle 1 (Bullet)","Vehicle 2 (Target)","Impact Velocity:"]

# slots=[Vehicle Type,Model Type,Year of Manufacture,Mass,Initial Velocity,Change in Velocity,Average Acceleration]
with open ("extract.txt", "r") as myfile:
            root=myfile.read().replace('\n', '')

#eed to change default font style and size
def first_page(table,slots,name):
	cells=table.columns[1].cells
	p=cells[0].add_paragraph(name)
	vehicle(1,cells[1],slots)
	vehicle(2,cells[2],slots)
	iv(cells[3],slots)
	mass(1,cells[4],slots)
	mass(2,cells[5],slots)
	mass_eq(cells[6],slots)
	'''velocity_change(1,table,slots)
	velocity_change(2,table,slots)
	average_acceleration(1,table,slots)
	average_acceleration(2,table,slots)'''

def n_page(n,document,slots):
	#document.add_page_break()
	h=document.add_heading(page1[n-1],0)
	h.alignment = WD_ALIGN_PARAGRAPH.CENTER
	table=document.add_table(rows=5,cols=2)
	cells=table.columns[0].cells
	cells[0].add_paragraph("Vehicle Type:").bold=True
	cells[1].add_paragraph("Year of Manufacture:").bold=True
	cells[2].add_paragraph("Model Type:").bold=True
	cells=table.columns[1].cells
	vehicle_type(n,cells[0],slots)
	year(n,cells[1],slots)
	vehicle_model(n,cells[2],slots)
	
	cells=table.columns[0].cells
	cells[3].add_paragraph("Bumper Type (Front):").bold=True
	cells[4].add_paragraph("Bumper Type (Rear):").bold=True
	p= document.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p.add_run("Before \n After \n After(Close Up)\n")
	table=document.add_table(rows=4,cols=2)
	cells=table.columns[0].cells
	cells[0].add_paragraph("Visible Damage (Y/N):").bold=True
	cells[1].add_paragraph("Repair Bill (Y/N):").bold=True
	cells[2].add_paragraph("Detail Visible Damage:").bold=True
	cells[3].add_paragraph("Repair Bill Assessment:").bold=True

'''def velocity_change(num,table,slots):
	cells=table.rows[num-1].cells
	string="Velocity Change "+str(num)+":"
	cells[0].text(string.ljust(35)).bold=True
	cells[1].text(slots[num-1][5])'''

'''def average_acceleration(num,table,slots):
	p = document.add_paragraph()
	string="Average Acceleration "+str(num)+":"
	p.add_run(string.ljust(35)).bold=True
	p.add_run(slots[num-1][6])'''

	
def vehicle(num,cell,slots):
	vehicle=slots[num-1][0]+ " "+slots[num-1][1]
	print vehicle
	p=cell.add_paragraph(vehicle)
	

def vehicle_type(num,cell,slots):
	new=slots[num-1][0]
	p=cell.add_paragraph(new)
	

def vehicle_model(num,cell,slots):
	new=slots[num-1][1]
	p=cell.add_paragraph(new)
	

def year(num,cell,slots):
	new=slots[num-1][2]
	p=cell.add_paragraph(new)
	
	
def iv(cell,slots):
	new=slots[0][4].replace("km/h","")
	p=cell.add_paragraph(new)
	
	
def mass(num,cell,slots):
	p=cell.add_paragraph(slots[num-1][3].replace("kg",""))
	
	
'''def sub(paragraph,string):
	r=paragraph.add_run()
	f=r.font
	f.subscript=True
	r.add_text(string).bold=True
	f.subscript=False'''
	
def mass_eq(cell,slots):
	m1=slots[0][3]
	m2=slots[1][3]
	m1=m1.replace("kg","")
	m2=m2.replace("kg","")
	m1=int(m1)
	m2=int(m2)
	meq=(m1*m2)/(m1+m2)
	p=cell.add_paragraph(str(meq))
	
	
start=0
for i in range(0,root.count("title\">Versuch")):
	slots=[]
	start= root.find("title\">Versuch ",start) + len("title\">Versuch ")
	end= root.find(":",start)
	test_name=root[start:end]
	print test_name
	for k in range (2):
		slots.append([])
		start=root.find("<tr>",start)
		for j in range(7):
			start= root.find("class=",start)
			start = root.find(">",start)+1
			end = root.find("<",start)
			slots[k].append(root[start:end])
	print slots
	document=Document()
	table=document.tables
	#document.add_page_break()
	#h=document.add_heading(test_name, 0)
	#h.alignment = WD_ALIGN_PARAGRAPH.CENTER		
	first_page(table[0],slots,test_name)
	n_page(1,document,slots)
	n_page(2,document,slots)
	document.add_page_break()		
	h=document.add_heading("Graphs", 0)
	h.alignment = WD_ALIGN_PARAGRAPH.CENTER
	print test_name+".docx"
	document.save("/home/tyron/DWA/Word_Documents/"+test_name+".docx")
'''	
document.add_heading('Heading, level 1', level=1)


#document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

document.add_page_break()

document.save('demo.docx') '''
